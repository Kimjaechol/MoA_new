"""
내용증명 작성 스킬 (Content Certification Writer)

법률적 의사표시 및 통지를 위한 내용증명 자동 작성 시스템

Author: LawPro AI Platform
Version: 1.0.0
License: MIT
"""

import os
import re
from dataclasses import dataclass
from typing import Dict, List, Optional, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@dataclass
class Party:
    """당사자 정보"""
    name: str  # 성명
    address: str  # 주소
    phone: Optional[str] = None  # 연락처
    email: Optional[str] = None  # 이메일
    registration_number: Optional[str] = None  # 주민등록번호/사업자번호


@dataclass
class ContentCertificationData:
    """내용증명 데이터"""
    type: str  # 내용증명 유형
    sender: Party  # 발신인
    receiver: Party  # 수신인
    title: str  # 제목
    facts: Dict[str, Any]  # 사실관계
    legal_basis: Optional[str] = None  # 법적 근거
    demand: Dict[str, Any] = None  # 요구사항
    deadline: Optional[str] = None  # 이행기한
    legal_action: Optional[str] = None  # 불이행 시 법적 조치
    attachments: Optional[List[str]] = None  # 첨부서류
    date: Optional[str] = None  # 작성일자


class ContentCertificationWriter:
    """내용증명 작성기"""

    # 내용증명 유형별 템플릿
    TEMPLATES = {
        "loan_repayment": "대여금 반환 청구",
        "lease_deposit_return": "임대차보증금 반환 청구",
        "unpaid_wages": "체불임금 청구",
        "retirement_pay": "퇴직금 청구",
        "product_payment": "물품대금 청구",
        "contract_termination": "계약 해제/해지 통지",
        "rent_demand": "임대료 독촉",
        "building_eviction": "건물명도 청구",
        "damage_compensation": "손해배상 청구",
        "property_transfer": "소유권이전등기 청구",
        "unfair_dismissal": "부당해고 이의 제기",
        "general": "일반 통지"
    }

    def __init__(self):
        """내용증명 작성기 초기화"""
        self.output_dir = "내용증명_출력"
        os.makedirs(self.output_dir, exist_ok=True)

    def create_document(self, data: Dict) -> Dict[str, str]:
        """
        내용증명 문서 생성

        Args:
            data: 내용증명 작성 데이터

        Returns:
            생성된 파일 정보
        """
        # 데이터 파싱
        cert_data = self._parse_data(data)

        # DOCX 생성
        doc = Document()
        self._set_document_style(doc)

        # 내용 작성
        self._write_title(doc, cert_data)
        self._write_parties(doc, cert_data)
        self._write_main_content(doc, cert_data)
        self._write_signature(doc, cert_data)

        # 파일 저장
        filename = self._generate_filename(cert_data)
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)

        return {
            "status": "success",
            "filename": filename,
            "filepath": filepath,
            "type": cert_data.type,
            "title": cert_data.title
        }

    def _parse_data(self, data: Dict) -> ContentCertificationData:
        """데이터 파싱"""
        sender_data = data.get("sender", {})
        receiver_data = data.get("receiver", {})

        sender = Party(
            name=sender_data.get("name"),
            address=sender_data.get("address"),
            phone=sender_data.get("phone"),
            email=sender_data.get("email"),
            registration_number=sender_data.get("registration_number")
        )

        receiver = Party(
            name=receiver_data.get("name"),
            address=receiver_data.get("address"),
            phone=receiver_data.get("phone"),
            email=receiver_data.get("email"),
            registration_number=receiver_data.get("registration_number")
        )

        cert_type = data.get("type", "general")
        title = data.get("title", self.TEMPLATES.get(cert_type, "통지서"))

        return ContentCertificationData(
            type=cert_type,
            sender=sender,
            receiver=receiver,
            title=title,
            facts=data.get("facts", {}),
            legal_basis=data.get("legal_basis"),
            demand=data.get("demand", {}),
            deadline=data.get("deadline"),
            legal_action=data.get("legal_action"),
            attachments=data.get("attachments", []),
            date=data.get("date", datetime.now().strftime("%Y년 %m월 %d일"))
        )

    def _set_document_style(self, doc: Document):
        """문서 스타일 설정"""
        # 페이지 설정
        section = doc.sections[0]
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # 기본 폰트 설정
        style = doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(11)

        # 한글 폰트 설정
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    def _write_title(self, doc: Document, data: ContentCertificationData):
        """제목 작성"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(data.title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = '맑은 고딕'

        doc.add_paragraph()  # 빈 줄

    def _write_parties(self, doc: Document, data: ContentCertificationData):
        """당사자 정보 작성"""
        # 발신인
        p = doc.add_paragraph()
        run = p.add_run("발신인: ")
        run.bold = True
        p.add_run(data.sender.name)

        if data.sender.address:
            doc.add_paragraph(f"주   소: {data.sender.address}")
        if data.sender.phone:
            doc.add_paragraph(f"연락처: {data.sender.phone}")

        doc.add_paragraph()  # 빈 줄

        # 수신인
        p = doc.add_paragraph()
        run = p.add_run("수신인: ")
        run.bold = True
        p.add_run(data.receiver.name)

        if data.receiver.address:
            doc.add_paragraph(f"주   소: {data.receiver.address}")

        doc.add_paragraph()  # 빈 줄

    def _write_main_content(self, doc: Document, data: ContentCertificationData):
        """본문 작성"""
        # 사실관계
        p = doc.add_paragraph()
        run = p.add_run("1. 사실관계")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph()

        # 유형별 사실관계 작성
        if data.type == "loan_repayment":
            self._write_loan_repayment_facts(doc, data)
        elif data.type == "lease_deposit_return":
            self._write_lease_deposit_facts(doc, data)
        elif data.type == "unpaid_wages":
            self._write_unpaid_wages_facts(doc, data)
        elif data.type == "retirement_pay":
            self._write_retirement_pay_facts(doc, data)
        elif data.type == "contract_termination":
            self._write_contract_termination_facts(doc, data)
        else:
            self._write_general_facts(doc, data)

        doc.add_paragraph()

        # 법적 근거 (있는 경우)
        if data.legal_basis:
            p = doc.add_paragraph()
            run = p.add_run("2. 법적 근거")
            run.bold = True
            run.font.size = Pt(12)

            doc.add_paragraph(data.legal_basis)
            doc.add_paragraph()

        # 요구사항
        p = doc.add_paragraph()
        run = p.add_run("3. 요구사항" if data.legal_basis else "2. 요구사항")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph()
        self._write_demand(doc, data)

        doc.add_paragraph()

        # 법적 조치 경고
        if data.legal_action:
            doc.add_paragraph(data.legal_action)
            doc.add_paragraph()

        # 첨부서류
        if data.attachments:
            p = doc.add_paragraph()
            run = p.add_run("첨부서류")
            run.bold = True

            for idx, attachment in enumerate(data.attachments, 1):
                doc.add_paragraph(f"{idx}. {attachment}", style='List Number')

    def _write_loan_repayment_facts(self, doc: Document, data: ContentCertificationData):
        """대여금 반환 사실관계 작성"""
        facts = data.facts

        paragraphs = [
            f"발신인은 {facts.get('loan_date', '____년 __월 __일')}에 수신인에게 {facts.get('loan_amount', '금 ____원')}을(를) 대여하였습니다.",
            f"위 대여금의 변제기는 {facts.get('repayment_date', '____년 __월 __일')}로 약정하였으며,",
        ]

        if facts.get('interest_rate'):
            paragraphs.append(f"이자는 {facts.get('interest_rate')}로 약정하였습니다.")

        if facts.get('proof'):
            paragraphs.append(f"이에 대한 증거로 {facts.get('proof')}이(가) 있습니다.")

        paragraphs.append("그러나 수신인은 위 변제기가 도과하였음에도 불구하고 현재까지 위 대여금을 변제하지 않고 있습니다.")

        for para in paragraphs:
            doc.add_paragraph(para)

    def _write_lease_deposit_facts(self, doc: Document, data: ContentCertificationData):
        """임대차보증금 반환 사실관계 작성"""
        facts = data.facts

        paragraphs = [
            f"발신인은 {facts.get('lease_start', '____년 __월 __일')}부터 {facts.get('lease_end', '____년 __월 __일')}까지 수신인과 임대차계약을 체결하였습니다.",
            f"임대차 목적물: {facts.get('property', '________________________')}",
            f"보증금: {facts.get('deposit', '금 ____원')}",
        ]

        if facts.get('monthly_rent'):
            paragraphs.append(f"월 차임: {facts.get('monthly_rent')}")

        paragraphs.extend([
            f"발신인은 위 임대차계약 기간이 만료되어 {facts.get('handover_date', '____년 __월 __일')} 목적물을 수신인에게 인도하였습니다.",
            f"목적물은 {facts.get('condition', '원상복구')} 상태로 인도되었습니다.",
            "그러나 수신인은 현재까지 위 보증금을 반환하지 않고 있습니다."
        ])

        for para in paragraphs:
            doc.add_paragraph(para)

    def _write_unpaid_wages_facts(self, doc: Document, data: ContentCertificationData):
        """체불임금 사실관계 작성"""
        facts = data.facts

        paragraphs = [
            f"발신인은 {facts.get('employment_start', '____년 __월 __일')}부터 {facts.get('employment_end', '____년 __월 __일')}까지 수신인의 {facts.get('position', '직원')}으로 근무하였습니다.",
            f"월 급여: {facts.get('monthly_salary', '금 ____원')}",
            f"체불 기간: {facts.get('unpaid_period', '____년 __월부터 __월까지')}",
            f"체불 금액: {facts.get('unpaid_amount', '금 ____원')}",
            "근로기준법 제43조에 따라 사용자는 근로자에게 임금을 지급할 의무가 있으나, 수신인은 현재까지 위 임금을 지급하지 않고 있습니다."
        ]

        for para in paragraphs:
            doc.add_paragraph(para)

    def _write_retirement_pay_facts(self, doc: Document, data: ContentCertificationData):
        """퇴직금 사실관계 작성"""
        facts = data.facts

        paragraphs = [
            f"발신인은 {facts.get('employment_start', '____년 __월 __일')}부터 {facts.get('employment_end', '____년 __월 __일')}까지 수신인의 {facts.get('position', '직원')}으로 근무하였습니다.",
            f"근속기간: {facts.get('service_years', '__')}년",
            f"퇴직일: {facts.get('retirement_date', '____년 __월 __일')}",
            f"퇴직금: {facts.get('retirement_pay', '금 ____원')}",
            "근로자퇴직급여보장법 제8조에 따라 사용자는 퇴직 근로자에게 퇴직일로부터 14일 이내에 퇴직금을 지급할 의무가 있으나, 수신인은 현재까지 위 퇴직금을 지급하지 않고 있습니다."
        ]

        for para in paragraphs:
            doc.add_paragraph(para)

    def _write_contract_termination_facts(self, doc: Document, data: ContentCertificationData):
        """계약 해제/해지 사실관계 작성"""
        facts = data.facts

        paragraphs = [
            f"발신인과 수신인은 {facts.get('contract_date', '____년 __월 __일')} {facts.get('contract_type', '____')} 계약을 체결하였습니다.",
            f"계약 내용: {facts.get('contract_details', '________________________')}",
        ]

        if facts.get('breach_details'):
            paragraphs.append(f"그러나 수신인은 {facts.get('breach_details')}하여 계약상 의무를 위반하였습니다.")

        paragraphs.append(f"따라서 발신인은 민법 제544조에 따라 본 통지로써 위 계약을 {facts.get('termination_type', '해제')}하고자 합니다.")

        for para in paragraphs:
            doc.add_paragraph(para)

    def _write_general_facts(self, doc: Document, data: ContentCertificationData):
        """일반 사실관계 작성"""
        facts = data.facts

        if isinstance(facts, dict):
            for key, value in facts.items():
                if isinstance(value, str):
                    doc.add_paragraph(value)
        elif isinstance(facts, list):
            for fact in facts:
                doc.add_paragraph(fact)
        elif isinstance(facts, str):
            doc.add_paragraph(facts)

    def _write_demand(self, doc: Document, data: ContentCertificationData):
        """요구사항 작성"""
        demand = data.demand

        if isinstance(demand, dict):
            if demand.get('amount'):
                doc.add_paragraph(f"청구 금액: {demand.get('amount')}")

            if demand.get('action'):
                doc.add_paragraph(f"요구 사항: {demand.get('action')}")

            deadline = demand.get('deadline') or data.deadline
            if deadline:
                doc.add_paragraph(f"이행 기한: {deadline}")

        elif isinstance(demand, str):
            doc.add_paragraph(demand)

        else:
            # 유형별 기본 요구사항
            if data.type == "loan_repayment":
                doc.add_paragraph("위 대여금 및 약정 이자를 지급하여 주시기 바랍니다.")
            elif data.type == "lease_deposit_return":
                doc.add_paragraph("위 임대차보증금 전액을 반환하여 주시기 바랍니다.")
            elif data.type in ["unpaid_wages", "retirement_pay"]:
                doc.add_paragraph("위 체불 금액을 즉시 지급하여 주시기 바랍니다.")

    def _write_signature(self, doc: Document, data: ContentCertificationData):
        """서명 작성"""
        doc.add_paragraph()
        doc.add_paragraph()

        # 작성일자
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(data.date)

        doc.add_paragraph()

        # 발신인 서명
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"발신인  {data.sender.name}")
        run.font.size = Pt(12)
        p.add_run("  (인)")

        doc.add_paragraph()

        # 수신인
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"{data.receiver.name} 귀하")
        run.font.size = Pt(12)

    def _generate_filename(self, data: ContentCertificationData) -> str:
        """파일명 생성"""
        today = datetime.now().strftime("%Y%m%d")
        type_name = self.TEMPLATES.get(data.type, "내용증명")
        receiver_name = data.receiver.name.replace(" ", "")

        return f"내용증명_{type_name}_{receiver_name}_{today}.docx"


# 사용 예시
if __name__ == "__main__":
    writer = ContentCertificationWriter()

    # 대여금 반환 청구 예시
    result = writer.create_document({
        "type": "loan_repayment",
        "title": "통 지 서 (대여금 반환 청구)",
        "sender": {
            "name": "홍길동",
            "address": "서울특별시 강남구 테헤란로 123",
            "phone": "010-1234-5678"
        },
        "receiver": {
            "name": "김철수",
            "address": "서울특별시 송파구 올림픽로 456"
        },
        "facts": {
            "loan_date": "2024년 1월 15일",
            "loan_amount": "금 오천만원 (₩50,000,000)",
            "repayment_date": "2024년 6월 15일",
            "interest_rate": "연 10%",
            "proof": "차용증 1부"
        },
        "legal_basis": "민법 제598조(소비대차)에 따라 차주는 대주에게 차용물을 반환할 의무가 있습니다.",
        "demand": {
            "amount": "원금 50,000,000원 + 이자 2,500,000원 = 총 52,500,000원",
            "deadline": "본 통지서 수령 후 7일 이내"
        },
        "legal_action": "위 기한 내 변제하지 않을 경우 민사소송 및 가압류 등 법적 조치를 취할 것임을 알려드립니다.",
        "attachments": ["차용증 사본 1부", "입금증 사본 1부"]
    })

    print(f"✅ 내용증명 생성 완료: {result['filename']}")
